#!/usr/bin/env python3
"""
job_generate.py - Stage 3: tailored application package generation.

For a queued job, this:
  1. fetches the full job description from the ATS,
  2. calls the Anthropic API with the master resume + JD to produce tailored
     resume content and a cover letter (accuracy over keyword stuffing, the
     operator's voice, no em dashes),
  3. writes versioned DOCX + PDF for both into the standardized package folder,
  4. records files and folder in the database and returns the package.

The CLI (`job draft IDENT`) wires this in and flips state to 'drafted'.

This NEVER applies to anything. It only produces documents for human review.

Dependencies (laptop):
  pip install anthropic python-docx requests pyyaml
  PDF: needs LibreOffice (`soffice`) on PATH, or set JOB_PDF=off to skip PDFs.

Auth:
  export ANTHROPIC_API_KEY=sk-ant-...
"""

import json
import os
import re
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path

import yaml

import job_monitor  # reuse its single polite Session (User-Agent) for JD fetches
import fit
import gate
import llm

APPLICATIONS_ROOT = Path(
    os.environ.get("JOB_APPS_DIR", Path.home() / "job-applications")
).expanduser()


# --- Job description fetch ------------------------------------------------

def posting_endpoint(row):
    """The public per-posting endpoint for a row, or None when there is none.

    The single definition of these URLs. fetch_description() reads the JD from
    them and liveness.check() asks the same URL whether the posting still
    exists, so a second copy of this table would drift from the fetchers and
    the two would silently disagree about what they were looking at.

    Returns None for an ATS with no public endpoint (Avature, iCIMS, manual
    entries) and for a Workday row whose company is not a real "{host}/{site}"
    cxs slug with a dotted host, which is what a manually-added or
    vanity-domain Workday row looks like.
    """
    ats = row["ats"]
    ext = row["ext_id"]
    company = row["company"]
    if ats == "greenhouse":
        return f"https://boards-api.greenhouse.io/v1/boards/{company}/jobs/{ext}"
    if ats == "lever":
        return f"https://api.lever.co/v0/postings/{company}/{ext}"
    if ats == "ashby":
        # Ashby's per-posting endpoint 401s on public boards; the public
        # job-board list carries every posting, so the endpoint is the board.
        return f"https://api.ashbyhq.com/posting-api/job-board/{company}"
    if ats == "smartrecruiters":
        return f"https://api.smartrecruiters.com/v1/companies/{company}/postings/{ext}"
    if ats == "workday":
        host, sep, site = company.partition("/")
        if sep and site and "." in host:
            tenant = host.split(".")[0]
            return f"https://{host}/wday/cxs/{tenant}/{site}{ext}"
    return None


def fetch_description(row):
    """Pull the full JD text for a job row, by ATS.

    Greenhouse/Lever/Ashby/SmartRecruiters all expose the description on the
    public posting endpoint, and are always re-fetched so a re-draft picks up an
    edited JD (generate() caches the fetched JD back into the description
    column, so we must not read that cache for a supported ATS).

    A manually-added job on an unsupported ATS (Avature, iCIMS, Workday, etc.)
    has no public endpoint, so it falls back to the stored description column,
    which lets it still be drafted.
    """
    ats = row["ats"]
    ext = row["ext_id"]
    # One definition of the URLs, shared with liveness.check (see
    # posting_endpoint). None means no public endpoint, including the
    # manually-added or vanity-domain Workday row that must fall through to
    # the stored JD below instead of building an invalid URL.
    url = posting_endpoint(row)
    if url and ats == "greenhouse":
        data = job_monitor.SESSION.get(url, timeout=20).json()
        return _strip_html(data.get("content", ""))
    if url and ats == "lever":
        data = job_monitor.SESSION.get(url, timeout=20).json()
        parts = [data.get("descriptionPlain", "")]
        for lst in data.get("lists", []):
            # In Lever's payload `text` is the section HEADING ("What You'll
            # Bring") and `content` is the requirement bullets under it. Taking
            # only `text` kept the headings and dropped every requirement, so a
            # 7k-char posting arrived as the company intro plus five labels
            # (measured: 932 chars stored for the Ellevation CloudOps role).
            head = _strip_html(lst.get("text", ""))
            body = _strip_html(lst.get("content", ""))
            parts.append(f"{head}\n{body}".strip() if head else body)
        return "\n\n".join(p for p in parts if p)
    if url and ats == "ashby":
        # The board list carries every posting's full description, so fetch it
        # and match on id (mirrors job_monitor.fetch_ashby).
        data = job_monitor.SESSION.get(url, timeout=20).json()
        for j in data.get("jobs", []):
            if str(j.get("id")) == str(ext):
                return j.get("descriptionPlain") or _strip_html(j.get("descriptionHtml", ""))
        return ""
    if url and ats == "smartrecruiters":
        data = job_monitor.SESSION.get(url, timeout=20).json()
        ad = (data.get("jobAd") or {}).get("sections", {})
        chunks = []
        for key in ("jobDescription", "qualifications", "additionalInformation"):
            sec = ad.get(key) or {}
            chunks.append(_strip_html(sec.get("text", "")))
        return "\n\n".join(c for c in chunks if c)
    if url and ats == "workday":
        data = job_monitor.SESSION.get(url, timeout=20).json()
        info = data.get("jobPostingInfo") or {}
        return _strip_html(info.get("jobDescription", ""))
    # Unsupported ATS (manual/Avature/iCIMS/Workday-vanity): no public endpoint,
    # so use the JD stored on the row if present.
    try:
        stored = row["description"]
    except (KeyError, IndexError):
        stored = None
    if stored:
        return stored
    raise ValueError(f"no description fetcher for ats '{ats}' and no stored description")


def _strip_html(html):
    if not html:
        return ""
    text = re.sub(r"<\s*br\s*/?>", "\n", html, flags=re.I)
    text = re.sub(r"</\s*(p|li|div|h\d)\s*>", "\n", text, flags=re.I)
    text = re.sub(r"<li[^>]*>", "- ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = (text.replace("&amp;", "&").replace("&lt;", "<")
                .replace("&gt;", ">").replace("&nbsp;", " ").replace("&#39;", "'"))
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --- Model call -----------------------------------------------------------

GEN_SYSTEM = """You tailor job application materials for one candidate.
You SELECT and EMPHASIZE from the candidate's master resume to fit a specific
job. You never invent experience, employers, dates, metrics, or skills that
are not in the master. If the job wants something the candidate lacks, you
omit it rather than fabricate.

Always include the candidate's AI experience, even when the job description
does not mention AI. Surface their AI and AI-accelerated DevOps work exactly as
it appears in the master resume (for example the AI-assisted operations
"wingman" for incident resolution, multi-model coding and review workflows,
multi-agent workflows, MCP-connected tooling, and AI agent rules that keep
AI-generated code on-convention) in the summary, in the skills, and in at least
one experience bullet. The candidate treats AI fluency as a baseline differentiator
for every role. Use only AI facts present in the master resume; do not invent
any, and do not add AI claims the master does not state.

Do not volunteer the candidate's gaps or weaknesses in the cover letter. Keep
the honest accounting of what the job wants but the candidate lacks in the
tailoring_note only. The cover letter stays confident and focuses on relevant
strengths; it never apologizes for or points out missing skills.

Hard voice rules:
- Plain English, direct, no filler.
- NEVER use em dashes or double hyphens. Use commas, parentheses, or separate
  sentences instead.
- Use consistent tense: present tense only for current roles and current facts,
  past tense for completed roles and completed achievements.
- Accuracy over keyword stuffing.

Return ONLY valid JSON, no markdown fences, with this exact shape:
{
  "company_name": "the target company's name, properly capitalized (e.g. GitLab, not gitlab)",
  "summary": "2-3 sentence tailored professional summary",
  "experience": [
    {"company": "...", "title": "...", "dates": "...", "location": "...",
     "points": ["tailored bullet", "..."]}
  ],
  "skills": ["grouped or flat skill strings most relevant to this job"],
  "cover_letter": "full cover letter body, plain paragraphs separated by blank lines",
  "tailoring_note": "1-2 sentences on what you emphasized and why, for the candidate to sanity-check"
}
Keep all experience entries from the master but reorder/trim bullets for fit.
"""


def _call_model(system, user, api_key):
    prov = llm.resolve_provider(component="draft")
    if api_key and prov.name == llm.DEFAULT_PROVIDER:
        prov = prov._replace(api_key=api_key)
    text = llm.call_messages(system, user, max_tokens=4000, provider=prov,
                             component="draft").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(json)?", "", text).rsplit("```", 1)[0].strip()
    return json.loads(text)


def call_model(master, jd_text, job_row, api_key):
    user = f"""MASTER RESUME (source of truth, do not exceed it):
{json.dumps(master, indent=2)}

TARGET JOB:
Company: {job_row['company']}
Title: {job_row['title']}
Location: {job_row['location']}

JOB DESCRIPTION:
{jd_text[:12000]}

Produce the tailored JSON now."""

    return _call_model(GEN_SYSTEM, user, api_key)


# --- Document writing -----------------------------------------------------

def _no_dashes(s):
    # Safety net: enforce the no-em-dash rule even if the model slips.
    return (s or "").replace("\u2014", ", ").replace("--", ", ")


def write_resume_docx(content, master, path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    c = master["contact"]
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(c["name"]); r.bold = True; r.font.size = Pt(18)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{c['location']} | {c['email']} | {c['phone']}").font.size = Pt(9.5)

    def heading(t):
        p = doc.add_paragraph()
        run = p.add_run(t.upper()); run.bold = True; run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    heading("Summary")
    doc.add_paragraph(_no_dashes(content["summary"]))

    heading("Experience")
    for e in content["experience"]:
        line = doc.add_paragraph()
        run = line.add_run(f"{e['title']}, {e['company']}"); run.bold = True
        meta = line.add_run(f"   {e.get('dates','')}  {e.get('location','')}")
        meta.italic = True; meta.font.size = Pt(9)
        for b in e["points"]:
            doc.add_paragraph(_no_dashes(b), style="List Bullet")

    heading("Skills")
    for s in content["skills"]:
        doc.add_paragraph(_no_dashes(s), style="List Bullet")

    if master.get("certifications"):
        heading("Certifications")
        for cert in master["certifications"]:
            doc.add_paragraph(_no_dashes(cert), style="List Bullet")

    if master.get("education"):
        heading("Education")
        for ed in master["education"]:
            doc.add_paragraph(f"{ed['degree']}, {ed['school']}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_cover_letter_docx(content, master, job_row, path):
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1.0)
        s.left_margin = s.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    c = master["contact"]
    doc.add_paragraph(c["name"])
    doc.add_paragraph(f"{c['email']} | {c['phone']}")
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph("")
    company_name = content.get("company_name") or job_row["company"]
    doc.add_paragraph(f"Re: {job_row['title']} at {company_name}")
    doc.add_paragraph("")
    for para in _no_dashes(content["cover_letter"]).split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def to_pdf(docx_path):
    """Convert a DOCX to PDF via LibreOffice if available. Returns path or None."""
    if os.environ.get("JOB_PDF", "").lower() == "off":
        return None
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir",
             str(docx_path.parent), str(docx_path)],
            check=True, capture_output=True, timeout=120,
        )
    except (subprocess.SubprocessError, OSError):
        return None
    pdf = docx_path.with_suffix(".pdf")
    return pdf if pdf.exists() else None


# --- Orchestration --------------------------------------------------------

def package_folder(job_row):
    slug = job_row["slug"]
    return APPLICATIONS_ROOT / f"{date.today().isoformat()}__{slug}"


def role_label(title):
    """'Engineering Manager, Cloud Safety' -> 'Engineering_Manager_Cloud_Safety'.

    Title_Case words, underscore-separated, punctuation removed. This is the
    recruiter-facing role string in the filename.
    """
    # Drop anything in parens, normalize separators to spaces.
    t = re.sub(r"\([^)]*\)", " ", title or "")
    t = re.sub(r"[^A-Za-z0-9]+", " ", t)
    words = [w for w in t.split() if w]
    # Title-case but preserve existing all-caps acronyms (SRE, AWS, ML).
    out = []
    for w in words:
        out.append(w if w.isupper() and len(w) <= 4 else w.capitalize())
    label = "_".join(out)
    return label[:80].strip("_") or "Role"


def applicant_label(master):
    """'Jordan Rivers' -> 'Jordan_Rivers'."""
    name = master.get("contact", {}).get("name", "Applicant")
    return re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")


def next_version(folder):
    """Next version number based on any resume files already in the folder."""
    vs = []
    for f in folder.glob("*_Resume_v*.docx"):
        m = re.search(r"_v(\d+)\.docx$", f.name)
        if m:
            vs.append(int(m.group(1)))
    return (max(vs) + 1) if vs else 1


def generate(db, job_row, master_path=None, api_key=None):
    """Full Stage 3 for one job. Returns a dict describing the package.

    Does not change job state; the CLI does that after a successful return.
    """
    # The Fit Gate. This is the choke point every artifact passes through, so
    # the CLI, the MCP tool a chat agent calls from Discord, and any future GUI all
    # inherit the block. Raises gate.GateBlocked.
    gate.require_pass(db, job_row)

    api_key = api_key or os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY not set")
    master_path = fit.resolve_config_path("master_resume.yaml", master_path)
    master = yaml.safe_load(master_path.read_text())

    jd = fetch_description(job_row)
    if not jd:
        raise RuntimeError(f"could not fetch a job description for {job_row['slug']}")

    content = call_model(master, jd, job_row, api_key)

    folder = package_folder(job_row)
    folder.mkdir(parents=True, exist_ok=True)
    v = next_version(folder)

    # Recruiter-facing filenames: Name_Role_DocType_vN. No company (the folder
    # carries that). Version suffix is strippable before upload.
    who = applicant_label(master)
    role = role_label(job_row["title"])
    base = f"{who}_{role}"

    resume_docx = folder / f"{base}_Resume_v{v}.docx"
    cover_docx = folder / f"{base}_Cover_Letter_v{v}.docx"
    write_resume_docx(content, master, resume_docx)
    write_cover_letter_docx(content, master, job_row, cover_docx)

    (folder / "job-description.md").write_text(
        f"# {job_row['title']} at {job_row['company']}\n\n"
        f"{job_row['url']}\n\n---\n\n{jd}\n")
    (folder / "tailoring-note.md").write_text(
        f"# Tailoring note (v{v})\n\n{content.get('tailoring_note','')}\n")

    written = {"resume_docx": resume_docx, "cover_docx": cover_docx}
    rp = to_pdf(resume_docx)
    cp = to_pdf(cover_docx)
    if rp:
        written["resume_pdf"] = rp
    if cp:
        written["cover_pdf"] = cp

    # Record in DB
    uid = job_row["uid"]
    db.set_fields(uid, folder=str(folder), description=jd)
    db.record_file(uid, "resume", str(resume_docx), version=v)
    if rp:
        db.record_file(uid, "resume_pdf", str(rp), version=v)
    db.record_file(uid, "cover_letter", str(cover_docx), version=v)
    if cp:
        db.record_file(uid, "cover_letter_pdf", str(cp), version=v)
    db.record_file(uid, "jd", str(folder / "job-description.md"), version=v)

    return {
        "folder": folder,
        "version": v,
        "files": written,
        "tailoring_note": content.get("tailoring_note", ""),
        "pdf": bool(rp and cp),
    }
